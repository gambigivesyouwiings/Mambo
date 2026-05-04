"""Convert FYP_Report_Draft.md to a Word document.

Lightweight markdown -> .docx converter tuned for the UoN FYP report style:
- ATX headings (# .. ######) -> Heading 1..6
- Tables (GitHub flavour) -> native Word tables
- Code fences -> mono-spaced paragraphs in a shaded box
- Inline code, bold, italics
- Bulleted and numbered lists (basic)
- Horizontal rules -> page breaks for top-level section transitions
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor


# ---------------------------------------------------------------------------
# Inline span parsing  (bold / italic / code)
# ---------------------------------------------------------------------------

INLINE_RE = re.compile(
    r"(\*\*([^*]+)\*\*|"     # 1,2: **bold**
    r"\*([^*]+)\*|"          # 3:    *italic*
    r"`([^`]+)`|"            # 4:    `code`
    r"\[([^\]]+)\]\(([^)]+)\))"  # 5,6: [text](link)
)


def _add_runs(paragraph, text: str):
    """Add inline-formatted runs for bold/italic/code/links to a paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        start, end = m.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        if m.group(2) is not None:
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3))
            r.italic = True
        elif m.group(4) is not None:
            r = paragraph.add_run(m.group(4))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif m.group(5) is not None:
            # [text](link) — render as italic blue
            r = paragraph.add_run(m.group(5))
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.italic = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    # |---|---|---| or | --- | :---: | etc.
    inner = s.strip("|").strip()
    cells = [c.strip() for c in inner.split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c)


def _parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _flush_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Code fence
# ---------------------------------------------------------------------------

def _add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Page numbering / header
# ---------------------------------------------------------------------------

def _add_page_number(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = 1  # center

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Main conversion driver
# ---------------------------------------------------------------------------

# Lines we want to drop from the markdown — pure formatting noise in .docx
SKIP_PATTERNS = (
    "<!--",
)

# Heading markers that should force a page break before the heading
PAGE_BREAK_HEADINGS = {
    "DECLARATION OF ORIGINALITY",
    "CERTIFICATION",
    "DEDICATION",
    "ACKNOWLEDGEMENTS",
    "TABLE OF CONTENTS",
    "LIST OF TABLES",
    "LIST OF FIGURES",
    "LIST OF ABBREVIATIONS",
    "ABSTRACT",
    "REFERENCES",
    "APPENDICES",
}


def _should_page_break(text: str, level: int) -> bool:
    if level == 1:
        return True
    if text.upper() in PAGE_BREAK_HEADINGS:
        return True
    # Top-level numbered sections: "1 INTRODUCTION", "2 LITERATURE REVIEW", ...
    if re.match(r"^\d+\s+[A-Z]", text) and level == 2:
        return True
    return False


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # Default normal paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Set 1-inch margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    _add_page_number(doc)

    # State machine
    i = 0
    n = len(lines)
    table_rows: list[list[str]] = []
    in_table = False
    in_code = False
    code_buffer: list[str] = []

    first_heading_seen = False

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # Skip blank lines (but flush any pending table)
        if not line.strip():
            if in_table:
                _flush_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue

        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(raw)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            if in_table:
                _flush_table(doc, table_rows)
                table_rows = []
                in_table = False
            # Treat as a soft section break — the page-break logic on
            # the next heading handles the hard break, so just skip.
            i += 1
            continue

        if any(line.lstrip().startswith(p) for p in SKIP_PATTERNS):
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|"):
            if (i + 1 < n) and _is_table_separator(lines[i + 1]) and not in_table:
                in_table = True
                table_rows = [_parse_table_row(line)]
                # Skip the separator line
                i += 2
                continue
            if in_table:
                if _is_table_separator(line):
                    i += 1
                    continue
                table_rows.append(_parse_table_row(line))
                i += 1
                continue
        else:
            if in_table:
                _flush_table(doc, table_rows)
                table_rows = []
                in_table = False

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            # Strip surrounding markdown emphasis
            heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
            if first_heading_seen and _should_page_break(heading_text, level):
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            first_heading_seen = True
            try:
                p = doc.add_heading("", level=min(level, 4))
            except Exception:
                p = doc.add_paragraph()
                p.style = doc.styles["Heading {}".format(min(level, 4))]
            _add_runs(p, heading_text)
            i += 1
            continue

        # Numbered list  ("1. ", "i. ", "a. ")
        m = re.match(r"^\s*([0-9]+|[a-z]|i|ii|iii|iv|v|vi|vii|viii|ix|x)[.)]\s+(.*)$", line, re.IGNORECASE)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(2))
            i += 1
            continue

        # Bullet list
        m = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # Block quote
        m = re.match(r"^>\s*(.*)$", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run("“" + m.group(1) + "”")
            r.italic = True
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_runs(p, line)
        i += 1

    if in_table:
        _flush_table(doc, table_rows)
    if in_code and code_buffer:
        _add_code_block(doc, code_buffer)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}  ({docx_path.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    md = Path(sys.argv[1] if len(sys.argv) > 1 else "FYP_Report_Draft.md")
    out = Path(sys.argv[2] if len(sys.argv) > 2 else md.with_suffix(".docx"))
    convert(md, out)
